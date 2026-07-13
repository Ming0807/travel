from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "manuals" / "ADMIN_USER_GUIDE_TH.docx"

NAVY = "123047"
TEAL = "0A6B62"
CORAL = "F3704C"
INK = "1F2937"
MUTED = "64748B"
LIGHT = "E8EEF5"
PALE_TEAL = "EAF6F3"
PALE_CORAL = "FFF1EC"
PALE_YELLOW = "FFF8E1"
WHITE = "FFFFFF"
RULE = "CBD5E1"


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color=RULE, size=5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=None, color=None, italic=None, latin="Calibri", thai="Leelawadee UI") -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), thai)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def style_paragraph_runs(paragraph, size=10.5, color=INK) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=8.5, color=MUTED)


def set_running_header_footer(section) -> None:
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(2)
    run = hp.add_run("SOUTHERN BORDER TOURISM  |  ADMIN HANDBOOK")
    set_run_font(run, size=8, bold=True, color=MUTED)
    p_pr = hp._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), RULE)
    pbdr.append(bottom)
    p_pr.append(pbdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("คู่มือการใช้งานระบบหลังบ้าน  |  หน้า ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(fp)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    set_running_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Heading 1", 16, NAVY, 18, 9),
        ("Heading 2", 13, TEAL, 13, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)

    doc.core_properties.title = "คู่มือการใช้งานระบบหลังบ้าน Southern Border Tourism"
    doc.core_properties.subject = "คู่มือผู้ดูแลระบบสำหรับอาจารย์และผู้ปฏิบัติงาน"
    doc.core_properties.author = "Southern Border Tourism Data & Intelligence Platform"
    doc.core_properties.keywords = "admin, CMS, tourism, dashboard, QR, คู่มือ"
    return doc


def add_para(doc, text="", *, bold_prefix=None, align=None, color=INK, size=10.5, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=size, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color)
    return p


def add_bullets(doc, items, *, numbered=False, level=0):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Inches(0.25 + 0.22 * level)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    style_paragraph_runs(p, size={1: 16, 2: 13, 3: 11.5}[level], color={1: NAVY, 2: TEAL, 3: NAVY}[level])
    return p


def add_callout(doc, title, body, *, tone="info"):
    colors = {
        "info": (LIGHT, NAVY),
        "success": (PALE_TEAL, TEAL),
        "warning": (PALE_YELLOW, "8A5A00"),
        "danger": (PALE_CORAL, "A63A25"),
    }
    fill, accent = colors[tone]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=accent, size=7)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 9300)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=170, bottom=120, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths is None:
        base = 9300 // len(headers)
        widths = [base] * len(headers)
        widths[-1] += 9300 - sum(widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r, size=9.3, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            if row_index % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=9.1, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_steps(doc, steps):
    for number, (title, body) in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.02)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(4)
        number_run = p.add_run(f"{number}. ")
        set_run_font(number_run, size=10.5, bold=True, color=CORAL)
        title_run = p.add_run(title)
        set_run_font(title_run, size=10.5, bold=True, color=NAVY)
        body_run = p.add_run(f" — {body}")
        set_run_font(body_run, size=10.5, color=INK)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
    kicker = add_para(
        doc,
        "คู่มือผู้ดูแลระบบ  |  ฉบับสำหรับการสอนและการปฏิบัติงาน",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=CORAL,
        size=10,
        after=18,
    )
    for run in kicker.runs:
        run.bold = True
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("คู่มือการใช้งาน\nระบบหลังบ้าน")
    set_run_font(r, size=30, bold=True, color=NAVY)
    subtitle = add_para(
        doc,
        "Southern Border Tourism Data & Intelligence Platform",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=TEAL,
        size=15,
        after=18,
    )
    for run in subtitle.runs:
        run.bold = True
    add_para(
        doc,
        "สำหรับจัดการข้อมูลท่องเที่ยว เนื้อหาหน้าบ้าน จุดเช็กอิน QR ใบประกาศดิจิทัล\nและข้อมูลเพื่อการวิเคราะห์การท่องเที่ยวชายแดนใต้",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=MUTED,
        size=11,
        after=30,
    )
    add_callout(
        doc,
        "หลักคิดของระบบ",
        "ระบบหลังบ้านไม่ได้มีไว้เพียงเพิ่มและแก้ไขข้อมูล แต่เป็นศูนย์ควบคุมคุณภาพข้อมูลที่นำไปใช้กับหน้าบ้าน การให้บริการนักท่องเที่ยว และแดชบอร์ดเพื่อการวางแผนอย่างยั่งยืน",
        tone="success",
    )
    add_para(
        doc,
        f"เวอร์ชันเอกสาร 1.0  |  จัดทำเมื่อ {date.today().strftime('%d/%m/%Y')}  |  เอกสารภายในโครงการ",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=MUTED,
        size=9,
        before=30,
        after=0,
    )
    add_page_break(doc)


def build_document():
    doc = setup_document()
    add_cover(doc)

    add_heading(doc, "เริ่มต้นใช้งานใน 5 นาที", 1)
    add_para(doc, "ส่วนนี้เหมาะสำหรับผู้ใช้ครั้งแรก เมื่อลงชื่อเข้าใช้สำเร็จแล้วจึงอ่านหัวข้อรายละเอียดตามงานที่รับผิดชอบ")
    add_heading(doc, "1. เข้าสู่ระบบ", 2)
    add_steps(doc, [
        ("เปิดหน้าเข้าสู่ระบบ", "เข้าเว็บไซต์ของโครงการ แล้วเติม /admin/login ต่อท้ายโดเมน เช่น https://ชื่อโดเมน/admin/login"),
        ("กรอกบัญชี", "กรอกชื่อผู้ใช้หรืออีเมล และรหัสผ่านที่ผู้ดูแลระบบมอบให้"),
        ("ตรวจสอบหน้าแรก", "เมื่อลงชื่อเข้าใช้สำเร็จ ระบบจะพาไปหน้า /admin และแสดงเมนูตามสิทธิของบัญชี"),
    ])
    add_callout(doc, "การรักษาความปลอดภัยบัญชี", "ไม่ใส่รหัสผ่านจริงไว้ในเอกสารนี้ ให้ส่งรหัสผ่านผ่านช่องทางแยกต่างหาก เปลี่ยนรหัสผ่านเมื่อรับบัญชี และไม่บันทึกรหัสผ่านบนเครื่องสาธารณะ", tone="danger")
    add_heading(doc, "2. วิธีอ่านหน้าจอหลังบ้าน", 2)
    add_table(doc, ["ส่วนของหน้าจอ", "หน้าที่", "วิธีใช้"], [
        ("เมนูด้านซ้าย", "เข้าโมดูลต่าง ๆ", "คลิกชื่อกลุ่มเพื่อขยาย แล้วเลือกหน้าที่ต้องการ"),
        ("แถบบน", "ดูบัญชีและเปิดหน้าบ้าน", "กด View Site เพื่อเปิดเว็บไซต์จริงในแท็บใหม่"),
        ("ตัวกรองและค้นหา", "ลดรายการให้เหลือเฉพาะที่ต้องการ", "ค้นก่อนแก้ไขเสมอ โดยเฉพาะเมื่อข้อมูลมีจำนวนมาก"),
        ("ป้ายสถานะ", "บอก Draft / Published / Active / Archived", "ตรวจทั้งสถานะเผยแพร่และสถานะเปิดใช้งาน"),
        ("ปุ่มบันทึก", "ยืนยันการเปลี่ยนแปลง", "รอข้อความบันทึกสำเร็จก่อนออกจากหน้า"),
    ], [1800, 2500, 5000])
    add_heading(doc, "3. ออกจากระบบ", 2)
    add_para(doc, "ระบบรุ่นปัจจุบันมีบริการออกจากระบบอยู่แล้ว แต่ปุ่มอาจยังไม่แสดงชัดในทุกหน้าจอ หากไม่พบปุ่มออกจากระบบ ให้หลีกเลี่ยงเครื่องสาธารณะ ปิดหน้าต่างส่วนตัวเมื่อใช้งานเสร็จ และแจ้งผู้ดูแลระบบเพื่อยุติเซสชัน บัญชีผู้ดูแลควรมีปุ่มออกจากระบบที่มองเห็นได้ก่อนส่งมอบใช้งานจริง")
    add_callout(doc, "ข้อควรจำ", "สิทธิของแต่ละบัญชีอาจไม่เหมือนกัน หากมองไม่เห็นเมนูหรือปุ่มใด ให้ตรวจ Role และ Permission ก่อนสรุปว่าระบบเสีย", tone="warning")

    add_page_break(doc)
    add_heading(doc, "แผนที่เมนูระบบหลังบ้าน", 1)
    add_para(doc, "ชื่อเมนูในระบบบางส่วนเป็นภาษาอังกฤษ ตารางนี้สรุปความหมายและงานหลักที่ควรใช้")
    add_table(doc, ["กลุ่ม", "เมนู", "ใช้ทำอะไร"], [
        ("Overview", "Home / Analytics", "ภาพรวมงานและแดชบอร์ดวิเคราะห์"),
        ("CRM", "Tourists", "ค้นหาและตรวจข้อมูลนักท่องเที่ยวแบบจำกัดตามสิทธิ"),
        ("CRM", "Bookings (Visits)", "ตรวจบันทึกการเข้าชมที่เกิดจาก QR Check-in"),
        ("CRM", "Reviews / Surveys", "กลั่นกรองรีวิวและตรวจแบบสำรวจ"),
        ("Content", "Content Hub", "จุดเริ่มต้นสำหรับจัดการเนื้อหาทั้งระบบ"),
        ("Content", "Content Health", "ตรวจความพร้อมและปัญหาของเนื้อหาที่เผยแพร่"),
        ("Content", "Destinations", "จัดการสถานที่ท่องเที่ยวและหน้ารายละเอียด"),
        ("Content", "Travel Routes", "จัดเส้นทาง วัน และจุดแวะพัก"),
        ("Content", "Articles", "จัดการเรื่องราวและบทความ"),
        ("Content", "Restaurants / Accommodations", "จัดการร้านอาหารและที่พัก"),
        ("Content", "Media Library", "คลังภาพกลางและข้อมูลกำกับภาพ"),
        ("Content", "Photo Spots", "จัดจุดถ่ายภาพที่เชื่อมกับสถานที่"),
        ("Operations", "Cert Templates", "จัดแม่แบบใบประกาศดิจิทัล"),
        ("Operations", "Checkin Codes", "สร้างและจัดการ QR ของสถานที่หรือจุดถ่ายภาพ"),
        ("Operations", "Badges / Leaderboard", "จัดแรงจูงใจ คะแนน และอันดับ"),
        ("System", "Users / Roles", "จัดผู้ใช้หลังบ้านและสิทธิการทำงาน"),
        ("System", "Audit Logs", "ตรวจประวัติการทำรายการสำคัญ"),
        ("System", "Messages", "จัดการข้อความจากผู้ใช้งาน"),
        ("System", "Settings", "ตั้งค่าระดับเว็บไซต์ หน้าแรก SEO และการเปิดปิดฟีเจอร์"),
    ], [1350, 2500, 5450])

    add_page_break(doc)
    add_heading(doc, "เข้าใจความสัมพันธ์ของ CMS ก่อนแก้ไข", 1)
    add_callout(doc, "กฎสำคัญที่สุด", "แก้ข้อมูลที่แหล่งข้อมูลหลักเพียงจุดเดียว แล้วให้ส่วนอื่นเลือกข้อมูลนั้นไปแสดง ไม่สร้างข้อมูลซ้ำใน Settings หรือกรอก URL ภาพซ้ำหลายที่", tone="success")
    add_table(doc, ["ต้องการเปลี่ยนอะไร", "แก้ที่เมนู", "ผลที่เกิดขึ้น"], [
        ("ชื่อ เนื้อหา พิกัด และรายละเอียดสถานที่", "Destinations", "หน้า /attractions/[slug] และส่วนที่อ้างอิงสถานที่นั้น"),
        ("ภาพปกหรือแกลเลอรีของสถานที่", "Destinations > Visual Editor > รูปภาพ", "หน้ารายละเอียดและการ์ดที่ใช้ภาพของสถานที่"),
        ("เลือกสถานที่ยอดนิยมบนหน้าแรก", "Settings > หน้าแรก > สถานที่ยอดนิยม", "เลือกรายการจาก Destinations โดยไม่สร้างสถานที่ซ้ำ"),
        ("บทความหรือเรื่องราว", "Articles", "หน้า Stories และบทความที่เกี่ยวข้อง"),
        ("เลือกเรื่องราวหรือเส้นทางบนหน้าแรก", "Settings > หน้าแรก", "ควบคุมลำดับและจำนวนที่แสดง"),
        ("ไฟล์ภาพกลาง", "Media Library", "นำกลับมาเลือกใช้ใน CMS หลายโมดูล"),
        ("ข้อความ Hero, Footer, SEO, Feature toggle", "Settings", "ค่าระดับเว็บไซต์ ไม่ใช่เนื้อหาเฉพาะรายการ"),
    ], [2800, 3100, 3400])
    add_heading(doc, "สถานะที่ต้องแยกให้ออก", 2)
    add_bullets(doc, [
        "Draft / ยังไม่เผยแพร่: บันทึกในระบบแล้ว แต่ผู้ใช้ทั่วไปยังไม่เห็น",
        "Published / เผยแพร่แล้ว: อนุญาตให้แสดงต่อสาธารณะ",
        "Active / เปิดใช้งาน: รายการยังถูกใช้งานในระบบ",
        "Archived / เก็บถาวร: ไม่ลบข้อมูลจริง เหมาะกับสื่อหรือรายการที่เลิกใช้แล้ว",
    ])
    add_callout(doc, "ทำไมบันทึกแล้วแต่หน้าบ้านยังไม่เห็น", "ตรวจว่า Published และ Active ครบหรือไม่ จากนั้นรอแคชหน้าสาธารณะประมาณ 60 วินาทีแล้วรีเฟรช หากยังไม่แสดงให้ตรวจ Content Health และความสัมพันธ์ของข้อมูล", tone="warning")

    add_page_break(doc)
    add_heading(doc, "จัดการสถานที่ท่องเที่ยว", 1)
    add_para(doc, "เมนู Destinations เป็นแหล่งข้อมูลหลักของสถานที่ และเป็นงาน CMS ที่สำคัญที่สุด")
    add_heading(doc, "สร้างสถานที่ใหม่", 2)
    add_steps(doc, [
        ("เปิด Destinations", "ค้นหาชื่อก่อนเพื่อป้องกันข้อมูลซ้ำ"),
        ("กดเพิ่มสถานที่", "กรอกชื่อไทย ชื่ออังกฤษ จังหวัด ประเภท และคำอธิบายสั้น"),
        ("บันทึกเป็น Draft", "ระบบตั้งใจให้ตรวจข้อมูลและภาพก่อนเผยแพร่"),
        ("เปิด Visual Editor", "แก้แต่ละส่วนในตำแหน่งที่ใกล้เคียงหน้าบ้านจริง"),
        ("ตรวจ Content Health", "แก้คำเตือนเรื่องภาพ ข้อมูลสำคัญ และสถานะ"),
        ("เผยแพร่", "เปิด Published และ Active แล้วเปิดหน้าบ้านเพื่อตรวจอีกครั้ง"),
    ])
    add_heading(doc, "ส่วนต่าง ๆ ในหน้ารายละเอียด", 2)
    add_table(doc, ["ส่วน", "เนื้อหาที่ควรใส่", "คำแนะนำ"], [
        ("Overview", "เรื่องย่อ จุดเด่น และข้อมูลสำคัญ", "เขียนให้สแกนอ่านง่าย ไม่ซ้ำกับทุกส่วน"),
        ("Things to Do", "กิจกรรมที่ทำได้จริง", "ใช้รายการและระบุฤดูกาลหรือข้อจำกัด"),
        ("Food & Drink", "อาหารเด่นหรือร้านที่เชื่อมโยง", "เชื่อมข้อมูลร้านอาหารเมื่อมี แทนการคัดลอกซ้ำ"),
        ("Travel Tips", "เวลาเหมาะสม การแต่งกาย ความปลอดภัย", "เขียนเป็นคำแนะนำที่นักท่องเที่ยวทำตามได้"),
        ("How to Get There", "เส้นทาง พิกัด และการเดินทาง", "ตรวจพิกัดบนแผนที่ก่อนเผยแพร่"),
        ("Reviews Summary", "ข้อมูลสรุปจากรีวิวที่ผ่านการกลั่นกรอง", "ไม่แต่งคะแนนหรือรีวิวตัวอย่าง"),
        ("Recommended Articles", "บทความที่เกี่ยวข้อง", "เลือกจาก Articles ที่เผยแพร่แล้ว"),
    ], [1900, 3350, 4050])
    add_callout(doc, "ข้อมูลจริงเท่านั้น", "หน้าสาธารณะควรใช้ข้อมูลที่บันทึกจริง หรือแสดง Empty State เมื่อยังไม่มีข้อมูล ห้ามใช้ Mockup หรือข้อความตัวอย่างที่ทำให้ผู้ใช้งานเข้าใจผิด", tone="danger")

    add_page_break(doc)
    add_heading(doc, "จัดการรูปภาพอย่างถูกวิธี", 1)
    add_heading(doc, "ใช้ Media Library เมื่อใด", 2)
    add_para(doc, "Media Library เป็นคลังสินทรัพย์กลาง มีประโยชน์สำหรับค้นหา เลือกใช้ซ้ำ ตรวจแหล่งที่มา และเก็บประวัติ ไม่ควรยุบหากระบบมีหลายโมดูลที่ใช้ภาพร่วมกัน")
    add_steps(doc, [
        ("เตรียมภาพ", "เลือกภาพคมชัด แนวนอนสำหรับปก และไม่มีลายน้ำที่ไม่ได้รับอนุญาต"),
        ("อัปโหลด", "อัปโหลด JPEG, PNG หรือ WebP ตามขนาดที่ระบบอนุญาต ระบบจะปรับภาพเป็น WebP และสร้าง thumbnail ในเส้นทางหลักของ Media Library"),
        ("กรอกข้อมูลกำกับ", "ระบุชื่อภาพ Alt text เครดิตผู้ถ่าย หมวดหมู่ และสิทธิการใช้งานเท่าที่มี"),
        ("เลือกจาก CMS", "กลับไปหน้าแก้ไขสถานที่หรือบทความ แล้วกดเลือกจาก Media Library"),
        ("บันทึกใน Drawer", "หลังเลือกภาพ ต้องกด บันทึกรูปภาพ ในแผงด้านข้างอีกครั้ง"),
        ("ตรวจหน้าบ้าน", "ดูทั้งมือถือและเดสก์ท็อปเพื่อเช็กการครอปและข้อความทับภาพ"),
    ])
    add_heading(doc, "มาตรฐานภาพที่แนะนำ", 2)
    add_table(doc, ["การใช้งาน", "สัดส่วนแนะนำ", "จุดตรวจ"], [
        ("ภาพปกสถานที่ / Hero", "16:9 หรือ 3:2", "วัตถุสำคัญอยู่กลางภาพ เผื่อพื้นที่ครอปบนมือถือ"),
        ("การ์ดรายการ", "4:3", "ยังเข้าใจภาพได้เมื่อย่อขนาด"),
        ("ภาพบทความ", "16:9", "สีและความสว่างสม่ำเสมอ"),
        ("ภาพใบประกาศ", "ตามแม่แบบแนวตั้งหรือแนวนอน", "ใบหน้าและข้อความไม่ถูกตัด"),
    ], [2800, 2300, 4200])
    add_bullets(doc, [
        "Alt text อธิบายสิ่งที่เห็นและประโยชน์ต่อบริบท ไม่ใช้คำว่า “รูปภาพ” เพียงอย่างเดียว",
        "ห้ามเก็บ URL ภายนอกแบบสุ่มเมื่อระบบมี Media Picker เพราะลิงก์อาจหมดอายุหรือโหลดช้า",
        "หากต้องเลิกใช้ภาพ ให้ Archive ก่อน ไม่ลบถาวรโดยไม่ตรวจว่าถูกใช้อยู่ที่ใด",
        "ภาพของนักท่องเที่ยวเป็นข้อมูลส่วนบุคคล ต้องจำกัดสิทธิและไม่ย้ายไปคลังสาธารณะโดยพลการ",
    ])

    add_page_break(doc)
    add_heading(doc, "ตั้งค่าหน้าแรกและเว็บไซต์", 1)
    add_para(doc, "Settings เหมาะกับค่าระดับเว็บไซต์ ไม่ใช่ CMS ชุดที่สอง เมื่อต้องการแก้เนื้อหาเฉพาะสถานที่ บทความ ร้านอาหาร หรือเส้นทาง ให้กลับไปแก้ที่โมดูลต้นทาง")
    add_table(doc, ["กลุ่ม Settings", "ควบคุมอะไร", "ข้อควรระวัง"], [
        ("ภาพหลักหน้าแรก (Hero)", "หัวเรื่อง คำอธิบาย ปุ่ม และภาพชุดแรก", "ตรวจทั้งมือถือและเดสก์ท็อป"),
        ("สถานที่ยอดนิยม", "เลือกสถานที่จาก Destinations และจัดลำดับ", "รายการต้อง Published และ Active"),
        ("เรื่องราวนักเดินทาง", "ชื่อส่วน คำอธิบาย ปุ่ม และจำนวนบทความ", "เนื้อหาจริงมาจาก Articles"),
        ("เส้นทางแนะนำ", "เลือกเส้นทางและเรียงลำดับ", "แก้รายละเอียดเส้นทางใน Travel Routes"),
        ("วิธีการทำงาน / ไฮไลต์ / CTA", "ข้อความอธิบายและภาพประกอบหน้าแรก", "บันทึกแล้วตรวจบริบททั้งหน้า"),
        ("หน้าสาธารณะ", "หัวและแบนเนอร์ของ Attractions, Stories, Routes, Restaurants", "ไม่เปลี่ยนจนสื่อความหมายผิดจากข้อมูลจริง"),
        ("SEO", "Title, Description และภาพแชร์เริ่มต้น", "ไม่ยัดคำค้น และไม่เผยข้อมูลภายใน"),
        ("Feature toggle", "เปิดหรือปิดโมดูลหลัก", "การปิดอาจกระทบเส้นทางผู้ใช้หลายหน้า"),
    ], [2700, 3600, 3000])
    add_heading(doc, "วิธีบันทึก Settings", 2)
    add_steps(doc, [
        ("เลือกกลุ่ม", "เปิดเฉพาะกลุ่มที่ต้องการแก้เพื่อลดความสับสน"),
        ("แก้ข้อมูล", "ระบบจะแสดงว่ามีรายการที่ยังไม่ได้บันทึก"),
        ("กดบันทึกทั้งหมด", "รอข้อความ “บันทึกสำเร็จ” และจำนวนรายการที่บันทึก"),
        ("เปิด View Site", "ตรวจหน้าแรกจริง โดยเฉพาะ Hero และส่วนที่เลือกรายการ"),
    ])
    add_callout(doc, "ตัวอย่าง: เปลี่ยนภาพสถานที่ยอดนิยม", "หากต้องการเปลี่ยนภาพของสถานที่ ให้แก้ภาพปกที่ Destinations ไม่ใช่ Settings แต่หากต้องการเปลี่ยนว่าสถานที่ใดขึ้นหน้าแรก ให้เลือกที่ Settings > สถานที่ยอดนิยม", tone="info")

    add_page_break(doc)
    add_heading(doc, "จัดการเรื่องราว เส้นทาง ร้านอาหาร และที่พัก", 1)
    add_table(doc, ["โมดูล", "ขั้นตอนหลัก", "สิ่งที่ต้องตรวจ"], [
        ("Articles", "สร้าง Draft > เขียนเนื้อหา > เลือกภาพปก > ตั้งค่า > Publish", "ผู้เขียน หมวดหมู่ จังหวัด Excerpt และภาพปก"),
        ("Travel Routes", "สร้างเส้นทาง > เพิ่มวันและจุดแวะพัก > จัดลำดับ > Publish", "ห้ามเพิ่มสถานที่ซ้ำโดยไม่ตั้งใจ และจำนวนวันต้องตรง"),
        ("Restaurants", "กรอกข้อมูลหลัก > เนื้อหา > พิกัด > ภาพ > Publish", "เวลาเปิด เมนูเด่น พิกัด และความสัมพันธ์กับสถานที่"),
        ("Accommodations", "กรอกข้อมูล > เลือกภาพ > ตรวจช่องทางติดต่อ > Publish", "ประเภทที่พัก ที่อยู่ พิกัด และข้อมูลติดต่อ"),
    ], [2000, 4400, 2900])
    add_heading(doc, "หลักการใช้ Visual Editor", 2)
    add_bullets(doc, [
        "คลิกส่วนที่ต้องการแก้บนหน้าจำลองของหน้า public",
        "แก้ข้อมูลใน Drawer ด้านข้าง แล้วกดปุ่มบันทึกของส่วนนั้น",
        "การเลือกภาพอย่างเดียวอาจเป็นเพียงสถานะชั่วคราว ต้องกด “บันทึกรูปภาพ” ก่อนปิด Drawer",
        "หากเห็นข้อความ “มีการเปลี่ยนแปลงที่ยังไม่ได้บันทึก” อย่ารีเฟรชหรือออกจากหน้า",
        "เปิดหน้าสาธารณะจริงหลังเผยแพร่เสมอ เพราะหน้าจำลองไม่แทนการทดสอบทุกขนาดหน้าจอ",
    ])
    add_callout(doc, "อย่าปิด Drawer เร็วเกินไป", "เมื่อเลือกภาพหรือแก้ข้อความแล้ว ให้กดบันทึกภายใน Drawer รอข้อความสำเร็จ แล้วจึงปิด หากรีเฟรชก่อนบันทึก การเปลี่ยนแปลงจะหาย", tone="warning")

    add_page_break(doc)
    add_heading(doc, "Photo Spots, QR Check-in และใบประกาศ", 1)
    add_para(doc, "QR เป็นกลไกบันทึกการเข้าชมและมอบคุณค่าให้นักท่องเที่ยว ไม่ใช่เป้าหมายหลักของฐานข้อมูล QR หนึ่งชุดควรเปิดเส้นทางกลาง /c/[code] และเชื่อมกับสถานที่หรือจุดถ่ายภาพที่ถูกต้อง")
    add_heading(doc, "สร้างจุดเช็กอินใหม่", 2)
    add_steps(doc, [
        ("ตรวจสถานที่", "สถานที่ต้องมีอยู่ใน Destinations และเปิดใช้งาน"),
        ("สร้าง Photo Spot", "ระบุชื่อ จุดตั้ง พิกัด และเชื่อม Attraction ที่ถูกต้อง"),
        ("สร้าง Checkin Code", "เลือก Attraction และ Photo Spot กำหนดรหัสที่อ่านง่าย สถานะ และช่วงเวลาใช้งาน"),
        ("บันทึกและดาวน์โหลด QR", "ใช้ปุ่มดาวน์โหลดจากรายการ Checkin Codes"),
        ("ทดสอบด้วยโทรศัพท์", "สแกน QR จริงผ่านกล้อง ตรวจว่าเปิด /c/[code] และแสดงชื่อสถานที่ถูกต้อง"),
        ("ทดสอบจนจบ", "ทำรายการแบบ Guest อัปโหลดภาพ สร้างใบประกาศ รับตราประทับ และเปิดแบบสำรวจต่อเนื่อง"),
    ])
    add_heading(doc, "เช็กลิสต์ก่อนนำ QR ไปติดตั้ง", 2)
    add_bullets(doc, [
        "ใช้โดเมน HTTPS ของโครงการและเส้นทาง /c/[code] เท่านั้น",
        "รหัส QR อยู่ในสถานะ Active และไม่หมดอายุ",
        "Attraction และ Photo Spot ที่เชื่อมยังเปิดใช้งาน",
        "พิมพ์ข้อความชื่อโครงการและ URL แบบตัวอักษรใต้ QR เพื่อให้ผู้ใช้ตรวจสอบก่อนเปิด",
        "ทดสอบทั้ง iPhone และ Android ผ่านเครือข่ายมือถือ ไม่ใช่เฉพาะ Wi-Fi",
        "เก็บทะเบียนตำแหน่งติดตั้งและผู้รับผิดชอบ เพื่อลดความเสี่ยง QR ปลอมถูกนำมาทับ",
    ])
    add_callout(doc, "ความปลอดภัยของ QR", "QR จริงควรมีตราหรือป้ายของโครงการ ระบุโดเมนปลายทาง และมีรอบตรวจสภาพหน้างาน ผู้ใช้ไม่ควรถูกบังคับให้ล็อกอินก่อนเห็นข้อมูลสถานที่หรือรับคุณค่าเบื้องต้น", tone="danger")
    add_heading(doc, "แม่แบบใบประกาศและตราประทับ", 2)
    add_para(doc, "ใช้ Cert Templates เพื่อกำหนดรูปแบบที่ระบบรองรับจริง ตรวจชื่อสถานที่ ภาพพื้นหลัง ตำแหน่งชื่อผู้รับ และการดาวน์โหลดบนมือถือ ส่วน Badges ใช้กำหนดเงื่อนไขและภาพรางวัล โดยตราประทับสถานที่ควรได้รับครั้งเดียวต่อคน แม้ Visits จะบันทึกการกลับมาเที่ยวได้หลายครั้ง")

    add_page_break(doc)
    add_heading(doc, "กลั่นกรองรีวิว ข้อความ และแบบสำรวจ", 1)
    add_heading(doc, "รีวิว", 2)
    add_steps(doc, [
        ("เปิด Reviews", "ค้นหาตามสถานที่ สถานะ หรือการเผยแพร่"),
        ("อ่านเนื้อหาและบริบท", "ตรวจถ้อยคำ ข้อมูลส่วนบุคคล สแปม และความเกี่ยวข้องกับสถานที่"),
        ("เลือกสถานะ", "เผยแพร่เฉพาะรีวิวที่ผ่านเกณฑ์ ไม่แก้คะแนนเพื่อให้ภาพรวมดูดีขึ้น"),
        ("จัดการอย่างย้อนกลับได้", "ใช้การซ่อนหรือสถานะที่เหมาะสมก่อนลบถาวร"),
    ])
    add_heading(doc, "Messages", 2)
    add_bullets(doc, [
        "จัดลำดับข้อความใหม่ ข้อร้องเรียน และคำถามที่กระทบความปลอดภัยก่อน",
        "อย่าคัดลอกข้อมูลติดต่อของผู้ใช้ไปไว้ในเอกสารหรือช่องทางที่ไม่ควบคุม",
        "บันทึกสถานะตอบกลับเพื่อให้ทีมไม่ตอบซ้ำหรือทำเรื่องตกหล่น",
    ])
    add_heading(doc, "Surveys", 2)
    add_para(doc, "แบบสำรวจเป็นข้อมูลติดตามหลังผู้ใช้ได้รับคุณค่าแล้ว ใช้เพื่อวิเคราะห์พฤติกรรม ค่าใช้จ่าย ความพึงพอใจ ความตั้งใจกลับมา และการแนะนำต่อ ไม่ควรตีความคำตอบเดี่ยวเป็นข้อสรุปของทั้งพื้นที่")
    add_callout(doc, "หลัก PDPA", "ผู้ดูแลควรเห็นข้อมูลส่วนบุคคลเท่าที่จำเป็นตามหน้าที่ ห้ามนำรูป ข้อมูลติดต่อ หรือประวัติการเดินทางไปใช้ภายนอกวัตถุประสงค์ที่แจ้งไว้", tone="danger")

    add_page_break(doc)
    add_heading(doc, "ตรวจข้อมูลนักท่องเที่ยวและการเข้าชม", 1)
    add_table(doc, ["เมนู", "ข้อมูลที่เห็น", "การตีความที่ถูกต้อง"], [
        ("Tourists", "โปรไฟล์นักท่องเที่ยวและอัตลักษณ์ที่เชื่อม", "หนึ่งคนอาจมี Guest, LINE, Email หรือ Google มากกว่าหนึ่งวิธี"),
        ("Bookings (Visits)", "การเข้าชมแต่ละครั้ง", "คนเดิมกลับมาได้หลายครั้ง จึงมีหลาย Visit"),
        ("Surveys", "คำตอบเชิงพฤติกรรม ค่าใช้จ่าย และความพึงพอใจ", "บาง Visit อาจไม่มี Survey เพราะเป็นข้อมูลติดตามแบบเลือกตอบ"),
        ("Badges / Passport", "รางวัลและตราประทับที่ได้รับ", "ตราสถานที่โดยทั่วไปไม่ควรออกซ้ำ แม้ Visit ซ้ำได้"),
    ], [1900, 3400, 4000])
    add_heading(doc, "แนวปฏิบัติสำหรับรายการข้อมูลจำนวนมาก", 2)
    add_bullets(doc, [
        "ใช้ตัวกรอง ค้นหา การเรียงลำดับ และ Pagination แทนการโหลดข้อมูลทั้งหมด",
        "ส่งออกเฉพาะคอลัมน์และช่วงเวลาที่จำเป็น",
        "ตรวจสิทธิ Export ก่อนดาวน์โหลดข้อมูลนักท่องเที่ยว",
        "ไม่นำไฟล์ Export ไปเก็บในอีเมลส่วนตัวหรือไดรฟ์สาธารณะ",
        "เมื่อไม่พบข้อมูล ให้แยก “ไม่มีข้อมูลตามตัวกรอง” ออกจาก “โหลดข้อมูลไม่สำเร็จ”",
    ])
    add_callout(doc, "ความหมายของข้อมูลประจำตัว", "Tourist profile เป็นระเบียนใช้งาน ไม่ใช่หลักฐานยืนยันบุคคลตามกฎหมาย และไม่ควรใช้เพื่อนับประชากรจริงโดยไม่ผ่านการทำความสะอาดและนิยามตัวชี้วัด", tone="warning")

    add_page_break(doc)
    add_heading(doc, "อ่าน Dashboard เพื่อการตัดสินใจ", 1)
    add_para(doc, "Dashboard มีหน้าที่ตอบคำถามด้านการวางแผน ไม่ใช่เพียงแสดงกราฟสวยงาม ก่อนนำตัวเลขไปใช้ควรตรวจนิยาม ช่วงเวลา ตัวกรอง และแหล่งข้อมูล")
    add_table(doc, ["หมวดวิเคราะห์", "คำถามที่ควรตอบ", "แหล่งข้อมูลหลัก"], [
        ("Executive overview", "ภาพรวมแนวโน้มสำคัญเป็นอย่างไร", "Visits, Tourists, Surveys"),
        ("Tourist profile", "นักท่องเที่ยวมาจากไหนและเป็นกลุ่มใด", "Tourists และ Geography"),
        ("Travel behavior", "เดินทางอย่างไร พักค้างหรือไม่ มากับใคร", "Surveys / Visit behavior"),
        ("Attraction performance", "สถานที่ใดได้รับความสนใจและพึงพอใจ", "Visits, Reviews, Satisfaction"),
        ("Expense analysis", "การใช้จ่ายกระจุกในหมวดหรือพื้นที่ใด", "Expense responses"),
        ("Satisfaction", "อะไรควรปรับปรุงและใครตั้งใจกลับมา", "Satisfaction surveys"),
        ("Sustainability", "พื้นที่ใดหนาแน่นหรือยังได้รับการส่งเสริมน้อย", "ข้อมูลรวมหลายมิติ"),
        ("Funnel", "ผู้ใช้หลุดที่ Scan, Check-in, Photo, Certificate หรือ Survey", "Events และ transaction records"),
    ], [2200, 4100, 3000])
    add_heading(doc, "ข้อควรระวังในการตีความ", 2)
    add_bullets(doc, [
        "QR Scan ไม่เท่ากับ Visit จนกว่ากระบวนการบันทึกการเข้าชมจะสำเร็จ",
        "จำนวน Tourist profile ไม่เท่ากับจำนวนคนจริง หากยังมีบัญชีซ้ำหรือ Guest หลายอุปกรณ์",
        "ค่าใช้จ่ายประมาณการไม่ใช่รายได้ที่ตรวจสอบทางบัญชีแล้ว",
        "ไม่มีข้อมูลไม่เท่ากับศูนย์ ต้องตรวจว่าผู้ใช้ยังไม่ได้ตอบหรือระบบโหลดล้มเหลว",
        "การเปรียบเทียบสถานที่ต้องพิจารณาช่วงเวลา จำนวนตัวอย่าง และความพร้อมของ QR",
    ])
    add_callout(doc, "ก่อนนำกราฟไปนำเสนอ", "บันทึกชื่อ Metric นิยาม ตารางต้นทาง ตัวกรอง สูตรคำนวณ และวันที่ดึงข้อมูลไว้ทุกครั้ง เพื่อให้ตรวจสอบย้อนกลับได้", tone="info")

    add_page_break(doc)
    add_heading(doc, "ผู้ใช้ บทบาท สิทธิ และ Audit Log", 1)
    add_heading(doc, "เพิ่มผู้ดูแลระบบ", 2)
    add_steps(doc, [
        ("กำหนดหน้าที่", "ระบุว่าบัญชีต้องทำงานกับ Content, Operations, Analytics หรือ System"),
        ("เลือก Role ที่เล็กที่สุด", "ให้เฉพาะสิทธิที่จำเป็นตามหลัก Least Privilege"),
        ("สร้างบัญชี", "ใช้ข้อมูลติดต่อของผู้ปฏิบัติงานจริงและส่งรหัสผ่านแยกช่องทาง"),
        ("ทดสอบสิทธิ", "ลงชื่อเข้าใช้ด้วยบัญชีทดสอบและตรวจทั้งเมนู ปุ่ม และ Server Action"),
        ("ทบทวนเป็นระยะ", "ปิดบัญชีที่ไม่ใช้งานและปรับสิทธิเมื่อหน้าที่เปลี่ยน"),
    ])
    add_heading(doc, "Audit Logs ใช้เมื่อใด", 2)
    add_bullets(doc, [
        "ตรวจว่าใครเปลี่ยนสถานะ เผยแพร่ Archive หรือแก้ค่าระบบ",
        "สืบหาสาเหตุเมื่อข้อมูลเปลี่ยนโดยไม่ทราบที่มา",
        "ประกอบการทบทวนสิทธิและเหตุการณ์ด้านความปลอดภัย",
        "Audit Log ไม่ควรบันทึกรหัสผ่าน คีย์ลับ หรือข้อมูลส่วนบุคคลเกินจำเป็น",
    ])
    add_callout(doc, "ห้ามใช้บัญชีร่วมกัน", "หากทุกคนใช้บัญชีเดียว Audit Log จะไม่สามารถระบุผู้ทำรายการจริง ควรสร้างบัญชีรายบุคคลและให้ Role ตามหน้าที่", tone="danger")

    add_page_break(doc)
    add_heading(doc, "คู่มือทำงานแบบรวดเร็ว", 1)
    add_heading(doc, "เปลี่ยนภาพสถานที่ยอดนิยม", 2)
    add_steps(doc, [
        ("เปิด Destinations", "ค้นหาสถานที่ที่ต้องการ"),
        ("เปิด Edit / Visual Editor", "เลือกส่วนรูปภาพหรือ Cover"),
        ("เลือกจาก Media Library", "ตรวจภาพตัวอย่างและข้อมูลกำกับ"),
        ("กดบันทึกรูปภาพ", "รอข้อความสำเร็จก่อนปิด Drawer"),
        ("ตรวจ Settings", "ตรวจว่าสถานที่ยังถูกเลือกใน “สถานที่ยอดนิยม”"),
        ("เปิดหน้าแรก", "รีเฟรชหลังแคชและตรวจมือถือกับเดสก์ท็อป"),
    ])
    add_heading(doc, "เพิ่มสถานที่และนำขึ้นหน้าแรก", 2)
    add_steps(doc, [
        ("สร้างสถานที่เป็น Draft", "กรอกข้อมูลหลักให้ครบ"),
        ("เพิ่มภาพและเนื้อหา", "ใช้ Visual Editor จนครบส่วนสำคัญ"),
        ("ตรวจ Content Health", "แก้ข้อผิดพลาดและคำเตือน"),
        ("Publish และ Active", "เปิดให้แสดงต่อสาธารณะ"),
        ("เลือกใน Settings", "ค้นหาและเพิ่มในสถานที่ยอดนิยม แล้วจัดลำดับ"),
        ("ตรวจหน้าบ้าน", "กดการ์ดและอ่านหน้ารายละเอียดจนจบ"),
    ])
    add_heading(doc, "สร้าง QR สำหรับสถานที่", 2)
    add_steps(doc, [
        ("สร้างหรือเลือก Photo Spot", "เชื่อมกับ Attraction ที่ถูกต้อง"),
        ("เพิ่ม Checkin Code", "ตั้งรหัส สถานะ และช่วงเวลา"),
        ("ดาวน์โหลด QR", "ใส่ในป้ายมาตรฐานของโครงการ"),
        ("สแกนทดสอบจริง", "ทำ flow จนได้ใบประกาศและตราประทับ"),
        ("บันทึกตำแหน่งติดตั้ง", "กำหนดผู้รับผิดชอบตรวจป้ายเป็นระยะ"),
    ])

    add_page_break(doc)
    add_heading(doc, "แก้ปัญหาที่พบบ่อย", 1)
    add_table(doc, ["อาการ", "ตรวจสอบตามลำดับ", "แนวทางแก้"], [
        ("มองไม่เห็นเมนูหรือปุ่ม", "Role > Permission > สถานะบัญชี", "ให้ Super Admin เพิ่มสิทธิที่จำเป็น ไม่แก้โค้ดเพื่อข้ามสิทธิ"),
        ("บันทึกแล้วแต่หน้าบ้านไม่เปลี่ยน", "ข้อความบันทึกสำเร็จ > Published > Active > Cache", "รอประมาณ 60 วินาที รีเฟรช และตรวจ Content Health"),
        ("เลือกรูปแล้วไม่ขึ้น", "เลือก Asset > ภาพ Preview > ปุ่มบันทึกรูปภาพ", "เปิด Drawer ใหม่ เลือกภาพ และกดบันทึกก่อนปิด"),
        ("ภาพแตกหรือถูกครอป", "สัดส่วน > จุดสำคัญในภาพ > ขนาดหน้าจอ", "เปลี่ยนภาพที่เหมาะกับสัดส่วนและตรวจมือถือ"),
        ("QR เปิดหน้าไม่พร้อมใช้งาน", "Code Active > วันหมดอายุ > Attraction/Photo Spot", "แก้สถานะหรือความสัมพันธ์ แล้วสแกนใหม่"),
        ("รายการค้นหาไม่เจอ", "ตัวกรอง > สถานะ > คำค้น > Pagination", "ล้างตัวกรองและค้นด้วยชื่อ/slug ที่สั้นลง"),
        ("Admin Section Error", "Environment > Migration > Permission > Server logs", "ส่ง Digest และเวลาที่เกิดเหตุให้ทีมพัฒนา ห้ามส่ง Secret"),
        ("Dashboard ไม่มีข้อมูล", "ช่วงเวลา > ตัวกรอง > Summary refresh > Source table", "แยก No data ออกจาก Query error ก่อนสรุปผล"),
    ], [2500, 3600, 3200])
    add_callout(doc, "ข้อมูลที่ควรส่งให้ทีมพัฒนา", "URL ที่เกิดปัญหา วันเวลา บัญชี/Role (ไม่ส่งรหัสผ่าน) ขั้นตอนก่อนเกิดเหตุ ภาพหน้าจอ ข้อความ Error และผลที่คาดหวัง", tone="info")

    add_page_break(doc)
    add_heading(doc, "เช็กลิสต์ก่อนเผยแพร่", 1)
    add_table(doc, ["ตรวจ", "รายการ", "ผ่านเมื่อ"], [
        ("□", "ชื่อไทย/อังกฤษและ Slug", "ถูกต้อง ไม่ซ้ำ และอ่านเข้าใจ"),
        ("□", "จังหวัด ประเภท พิกัด และความสัมพันธ์", "เชื่อมข้อมูลจริงครบ"),
        ("□", "ภาพปก แกลเลอรี Alt text และเครดิต", "โหลดได้ ครอปดี และมีสิทธิใช้"),
        ("□", "เนื้อหาแต่ละส่วน", "ไม่มีข้อความ Mockup หรือข้อมูลเก่า"),
        ("□", "Published และ Active", "ตั้งตรงตามวันที่ต้องการเผยแพร่"),
        ("□", "Content Health", "ไม่มีปัญหาระดับสูงที่ยังไม่ได้แก้"),
        ("□", "หน้าบ้านบนมือถือ", "ข้อความไม่ล้น ปุ่มกดได้ และภาพไม่บังเนื้อหา"),
        ("□", "หน้าบ้านบนเดสก์ท็อป", "ลำดับเนื้อหาและภาพสมดุล"),
        ("□", "ลิงก์และปุ่ม", "ทุกปุ่มไปยังหน้าที่มีอยู่จริง"),
        ("□", "ข้อมูลส่วนบุคคล", "ไม่เผยข้อมูลเกินวัตถุประสงค์หรือสิทธิ"),
    ], [750, 4000, 4550])
    add_heading(doc, "เช็กลิสต์สิ้นวันสำหรับผู้ดูแล", 2)
    add_bullets(doc, [
        "ไม่มีรายการค้างบันทึกใน Drawer หรือ Settings",
        "ตรวจข้อความ รีวิว หรือเหตุการณ์สำคัญที่ยังไม่จัดการ",
        "ไม่เหลือไฟล์ Export บนเครื่องสาธารณะ",
        "ปิดแท็บหลังบ้านและออกจากระบบเมื่อมีปุ่มรองรับ",
        "แจ้งปัญหาที่กระทบผู้ใช้หรือข้อมูลให้ผู้รับผิดชอบทันที",
    ])

    add_page_break(doc)
    add_heading(doc, "ขอบเขตความรับผิดชอบและช่องทางช่วยเหลือ", 1)
    add_table(doc, ["เหตุการณ์", "ผู้รับผิดชอบหลัก", "การดำเนินการ"], [
        ("แก้เนื้อหา ภาพ หรือสถานะเผยแพร่", "Content Admin", "ตรวจและบันทึกตามคู่มือนี้"),
        ("สิทธิไม่พอหรือบัญชีถูกปิด", "Super Admin", "ตรวจ Users, Roles และ Audit Logs"),
        ("QR หรือ Certificate ใช้งานไม่ได้", "Operations + Developer", "เก็บ URL/code/visit id และเวลาที่เกิดเหตุ"),
        ("ข้อมูล Dashboard ผิดปกติ", "Data/Admin + Developer", "ตรวจนิยาม ตัวกรอง และตารางต้นทาง"),
        ("สงสัยข้อมูลรั่วไหล", "Project Owner / Privacy Lead", "หยุดแชร์ข้อมูล เก็บหลักฐาน และแจ้งทันที"),
        ("Server, Deployment หรือ Database error", "Developer / DevOps", "แนบ log ที่ตัด Secret แล้วและ Digest"),
    ], [2650, 2650, 4000])
    add_callout(doc, "ข้อมูลโครงการ", "ชื่อระบบ: Southern Border Tourism Data & Intelligence Platform\nขอบเขตพื้นที่: ยะลา ปัตตานี และนราธิวาส\nวัตถุประสงค์: สร้างฐานข้อมูลท่องเที่ยวคุณภาพสูงเพื่อการวางแผนการท่องเที่ยวอย่างยั่งยืน", tone="success")
    add_heading(doc, "บันทึกการส่งมอบบัญชี", 2)
    add_table(doc, ["รายการ", "กรอกเมื่อส่งมอบ"], [
        ("ชื่อผู้ใช้งาน", "________________________________________"),
        ("Role", "________________________________________"),
        ("วันที่ส่งมอบ", "________________________________________"),
        ("ผู้ส่งมอบ", "________________________________________"),
        ("ช่องทางแจ้งปัญหา", "________________________________________"),
    ], [2800, 6500])
    add_para(doc, "หมายเหตุ: ส่งรหัสผ่านผ่านช่องทางแยกจากเอกสาร และไม่เขียนรหัสผ่านลงในแบบฟอร์มนี้", color="A63A25", size=9.5, before=6, after=0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
